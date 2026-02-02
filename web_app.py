from flask import Flask, jsonify, request, session
import io
import re
import statistics
import os
import tempfile
from docx import Document
import pdfplumber
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

from predict import get_detailed_prediction, predict_text

app = Flask(__name__, static_folder="frontend", static_url_path="")
app.secret_key = os.environ.get("SECRET_KEY", "your-secret-key-change-this")

# Google Drive OAuth Configuration
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
CLIENT_SECRETS_FILE = "client_secret.json"  # You'll create this from Google Cloud Console


@app.route("/")
def index():
    return app.send_static_file("index.html")


@app.post("/api/predict")
def api_predict():
    payload = request.get_json(silent=True) or {}
    text = (payload.get("text") or "").strip()

    if len(text) < 10:
        return jsonify(error="Please provide at least 10 characters."), 400

    try:
        result = get_detailed_prediction(text)
    except FileNotFoundError as exc:
        return jsonify(error=str(exc)), 500
    except Exception as exc:
        return jsonify(error=f"Prediction failed: {exc}"), 500

    safe_result = {
        "prediction": result.get("prediction"),
        "confidence": float(result.get("confidence", 0)),
        "human_probability": float(result.get("human_probability", 0)),
        "ai_probability": float(result.get("ai_probability", 0)),
        "word_count": int(result.get("word_count", 0)),
        "warning": result.get("warning"),
        "needs_review": bool(result.get("needs_review")),
    }

    return jsonify(safe_result)


@app.post("/api/extract")
def api_extract():
    if "file" not in request.files:
        return jsonify(error="No file uploaded."), 400

    file = request.files["file"]
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        with pdfplumber.open(file.stream) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        text = "\n".join(pages).strip()
    elif filename.endswith(".docx"):
        content = io.BytesIO(file.read())
        doc = Document(content)
        text = "\n".join([p.text for p in doc.paragraphs]).strip()
    elif filename.endswith((".txt", ".md", ".csv", ".json")):
        text = (file.read() or b"").decode("utf-8", errors="ignore").strip()
    else:
        return jsonify(error="Unsupported file type. Use .pdf, .docx, .txt, .md, .csv, or .json."), 400

    if not text:
        return jsonify(error="Unable to extract text from file."), 400

    return jsonify(text=text)


# Google Drive Authentication Routes
@app.route("/api/drive/auth")
def drive_auth():
    """Initiate Google Drive OAuth flow"""
    try:
        flow = Flow.from_client_secrets_file(
            CLIENT_SECRETS_FILE,
            scopes=SCOPES,
            redirect_uri=request.url_root + 'api/drive/callback'
        )
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true'
        )
        session['state'] = state
        return jsonify(auth_url=authorization_url)
    except FileNotFoundError:
        return jsonify(error="client_secret.json not found. Please configure Google OAuth."), 500
    except Exception as e:
        return jsonify(error=f"OAuth setup failed: {str(e)}"), 500


@app.route("/api/drive/callback")
def drive_callback():
    """Handle OAuth callback"""
    try:
        state = session.get('state')
        flow = Flow.from_client_secrets_file(
            CLIENT_SECRETS_FILE,
            scopes=SCOPES,
            state=state,
            redirect_uri=request.url_root + 'api/drive/callback'
        )
        flow.fetch_token(authorization_response=request.url)
        
        credentials = flow.credentials
        session['credentials'] = {
            'token': credentials.token,
            'refresh_token': credentials.refresh_token,
            'token_uri': credentials.token_uri,
            'client_id': credentials.client_id,
            'client_secret': credentials.client_secret,
            'scopes': credentials.scopes
        }
        
        return """
        <html>
            <body>
                <script>
                    window.opener.postMessage({type: 'drive-auth-success'}, '*');
                    window.close();
                </script>
                <p>Authentication successful! You can close this window.</p>
            </body>
        </html>
        """
    except Exception as e:
        return f"<html><body><p>Authentication failed: {str(e)}</p></body></html>", 400


@app.route("/api/drive/status")
def drive_status():
    """Check if user is authenticated with Google Drive"""
    if 'credentials' in session:
        return jsonify(authenticated=True)
    return jsonify(authenticated=False)


@app.post("/api/drive/download")
def drive_download():
    """Download file from Google Drive and extract text"""
    if 'credentials' not in session:
        return jsonify(error="Not authenticated with Google Drive"), 401
    
    payload = request.get_json(silent=True) or {}
    file_id = payload.get("file_id")
    
    if not file_id:
        return jsonify(error="No file_id provided"), 400
    
    try:
        credentials = Credentials(**session['credentials'])
        service = build('drive', 'v3', credentials=credentials)
        
        # Get file metadata
        file_metadata = service.files().get(fileId=file_id, fields='name,mimeType').execute()
        filename = file_metadata.get('name', '').lower()
        mime_type = file_metadata.get('mimeType', '')
        
        # Download file content
        request_obj = service.files().get_media(fileId=file_id)
        file_stream = io.BytesIO()
        downloader = MediaIoBaseDownload(file_stream, request_obj)
        
        done = False
        while not done:
            status, done = downloader.next_chunk()
        
        file_stream.seek(0)
        
        # Extract text based on file type
        text = ""
        
        if 'pdf' in mime_type or filename.endswith('.pdf'):
            with pdfplumber.open(file_stream) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            text = "\n".join(pages).strip()
            
        elif 'word' in mime_type or filename.endswith('.docx'):
            doc = Document(file_stream)
            text = "\n".join([p.text for p in doc.paragraphs]).strip()
            
        elif 'text' in mime_type or filename.endswith(('.txt', '.md')):
            text = file_stream.read().decode('utf-8', errors='ignore').strip()
            
        elif 'document' in mime_type:  # Google Docs
            # Export as plain text
            request_obj = service.files().export_media(fileId=file_id, mimeType='text/plain')
            file_stream = io.BytesIO()
            downloader = MediaIoBaseDownload(file_stream, request_obj)
            done = False
            while not done:
                status, done = downloader.next_chunk()
            file_stream.seek(0)
            text = file_stream.read().decode('utf-8', errors='ignore').strip()
            
        else:
            return jsonify(error=f"Unsupported file type: {mime_type}"), 400
        
        if not text:
            return jsonify(error="Unable to extract text from file"), 400
        
        return jsonify(text=text, filename=file_metadata.get('name'))
        
    except Exception as e:
        return jsonify(error=f"Failed to download file: {str(e)}"), 500


def _split_sentences(text):
    return [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]


def _compute_reasons(text):
    reasons = []
    sentences = _split_sentences(text)
    words = re.findall(r"\b\w+\b", text.lower())

    if sentences:
        lengths = [len(re.findall(r"\b\w+\b", s)) for s in sentences]
        if len(lengths) > 1:
            variance = statistics.pvariance(lengths)
            if variance < 8:
                reasons.append("Low sentence length variation")

    if words:
        bigrams = list(zip(words, words[1:]))
        if bigrams:
            counts = {}
            for bg in bigrams:
                counts[bg] = counts.get(bg, 0) + 1
            top = max(counts.values())
            if top / max(len(bigrams), 1) > 0.08:
                reasons.append("Repetitive phrasing patterns detected")

        function_words = {
            "the",
            "and",
            "to",
            "of",
            "in",
            "that",
            "for",
            "on",
            "with",
            "as",
            "is",
            "it",
            "this",
            "by",
            "from",
        }
        ratio = sum(1 for w in words if w in function_words) / max(len(words), 1)
        if ratio > 0.55:
            reasons.append("High function-word density")

    if not reasons:
        reasons.append("Balanced structure and phrasing")

    return reasons


@app.post("/api/analyze")
def api_analyze():
    payload = request.get_json(silent=True) or {}
    text = (payload.get("text") or "").strip()

    if len(text) < 10:
        return jsonify(error="Please provide at least 10 characters."), 400

    try:
        overall = get_detailed_prediction(text)
    except FileNotFoundError as exc:
        return jsonify(error=str(exc)), 500
    except Exception as exc:
        return jsonify(error=f"Analysis failed: {exc}"), 500

    sentences = _split_sentences(text)
    sentence_results = []
    for sentence in sentences:
        label, confidence, probs, warning, word_count = predict_text(sentence)
        sentence_results.append(
            {
                "text": sentence,
                "label": label,
                "confidence": float(confidence),
                "human_probability": float(probs[0] * 100),
                "ai_probability": float(probs[1] * 100),
                "word_count": int(word_count),
                "warning": warning,
            }
        )

    safe_overall = {
        "prediction": overall.get("prediction"),
        "confidence": float(overall.get("confidence", 0)),
        "human_probability": float(overall.get("human_probability", 0)),
        "ai_probability": float(overall.get("ai_probability", 0)),
        "word_count": int(overall.get("word_count", 0)),
        "warning": overall.get("warning"),
        "needs_review": bool(overall.get("needs_review")),
    }

    return jsonify({
        "overall": safe_overall,
        "sentences": sentence_results,
        "reasons": _compute_reasons(text),
    })
    


if __name__ == "__main__":
    app.run(host="127.0.0.1", port=8000, debug=True)